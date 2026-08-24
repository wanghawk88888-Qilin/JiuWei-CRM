"""Resume parser service — extracts raw text from docx and pdf files."""

import logging

logger = logging.getLogger(__name__)

# Below this many characters a document is treated as having no usable text.
# Scanned / image-only PDFs typically yield 0 characters, occasionally a few
# stray glyphs from a header stamp.
MIN_USABLE_TEXT_LENGTH = 10


class ResumeParseError(Exception):
    """Raised when a resume file cannot be turned into usable text.

    Carries a stable ``error_code`` so the API and the batch UI can explain
    *why* a file failed without parsing English exception text.
    """

    def __init__(self, error_code: str, message: str):
        super().__init__(f"{error_code}|{message}")
        self.error_code = error_code
        self.message = message


def parse_docx(file_path: str) -> str:
    """Extract text from a .docx file, including table cells.

    Many Chinese resumes lay the personal-info block out as a table, so table
    text matters as much as paragraph text.
    """
    from docx import Document

    doc = Document(file_path)
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Drop the duplicates python-docx reports for merged cells.
            deduped: list[str] = []
            for cell in cells:
                if cell and (not deduped or deduped[-1] != cell):
                    deduped.append(cell)
            if deduped:
                parts.append("  ".join(deduped))

    return "\n".join(parts)


def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file.

    Tries PyMuPDF first; falls back to pdfplumber if available. No OCR is
    attempted — image-only PDFs are reported as unparseable rather than
    guessed at.
    """
    try:
        import fitz  # PyMuPDF

        text_parts: list[str] = []
        with fitz.open(file_path) as doc:
            for page in doc:
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text.strip())
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("PyMuPDF not available, trying pdfplumber")

    try:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
        return "\n".join(text_parts)
    except ImportError:
        raise ResumeParseError(
            "PARSER_UNAVAILABLE", "服务器缺少 PDF 解析组件"
        ) from None


def parse_resume_text(file_path: str, file_type: str) -> str:
    """Parse a resume file and return its extracted text.

    Args:
        file_path: Path to the resume file.
        file_type: One of 'docx' or 'pdf' (leading dot optional).

    Returns:
        Extracted plain text.

    Raises:
        ResumeParseError: Unsupported type, unreadable file, or a PDF with no
            extractable text (``PDF_NO_EXTRACTABLE_TEXT``).
    """
    normalized_type = file_type.lower().lstrip(".")

    if normalized_type == "docx":
        try:
            text = parse_docx(file_path)
        except ResumeParseError:
            raise
        except Exception as exc:
            logger.warning("DOCX parse failed for %s: %s", file_path, exc)
            raise ResumeParseError(
                "DOCX_PARSE_FAILED", "Word 文件解析失败，可能已损坏"
            ) from exc

        if len(text.strip()) < MIN_USABLE_TEXT_LENGTH:
            raise ResumeParseError(
                "DOCX_NO_EXTRACTABLE_TEXT", "Word 文件中未提取到有效文本"
            )
        return text

    if normalized_type == "pdf":
        try:
            text = parse_pdf(file_path)
        except ResumeParseError:
            raise
        except Exception as exc:
            logger.warning("PDF parse failed for %s: %s", file_path, exc)
            raise ResumeParseError(
                "PDF_PARSE_FAILED", "PDF 文件解析失败，可能已损坏"
            ) from exc

        # Scanned / image-only PDFs land here. v0.2.1 deliberately does not
        # introduce an OCR dependency; the file goes to a human instead.
        if len(text.strip()) < MIN_USABLE_TEXT_LENGTH:
            raise ResumeParseError(
                "PDF_NO_EXTRACTABLE_TEXT",
                "PDF 无法提取文本（可能是扫描件/图片版），请人工处理",
            )
        return text

    raise ResumeParseError("INVALID_FILE_TYPE", f"不支持的文件类型: {file_type}")
